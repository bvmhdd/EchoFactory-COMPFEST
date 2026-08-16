import os
import shutil
import urllib.request
import base64
import zlib
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

WORKSPACE_DIR = r"c:\Users\muhib\Downloads\COMPFEST"
SCRATCH_DIR = os.path.join(WORKSPACE_DIR, "scratch_diagrams")
os.makedirs(SCRATCH_DIR, exist_ok=True)

TARGET_DOCX = os.path.join(WORKSPACE_DIR, "EchoFactory_Diagrams.docx")
BACKUP_DOCX = os.path.join(WORKSPACE_DIR, "EchoFactory_Diagrams_Backup.docx")

# 1. Backup file lama jika ada
if os.path.exists(TARGET_DOCX) and not os.path.exists(BACKUP_DOCX):
    shutil.copy2(TARGET_DOCX, BACKUP_DOCX)
    print(f"[OK] Backup dibuat: {BACKUP_DOCX}")

def render_diagram(code_text, diag_type, output_name):
    compressed = zlib.compress(code_text.encode("utf-8"), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode("ascii")
    url = f"https://kroki.io/{diag_type}/png/{encoded}"
    
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    out_path = os.path.join(SCRATCH_DIR, output_name)
    with urllib.request.urlopen(req, timeout=30) as resp:
        with open(out_path, "wb") as f:
            f.write(resp.read())
    print(f"[OK] Rendered {output_name} ({os.path.getsize(out_path)} bytes)")
    return out_path

# ==================== DEFINISI DIAGRAM ====================

USECASE_PUML = """@startuml
skinparam actorStyle awesome
skinparam packageStyle rectangle
skinparam roundcorner 12
skinparam shadowing false
skinparam defaultFontName Arial
skinparam defaultFontSize 12

skinparam actor {
    BackgroundColor #E8EAF6
    BorderColor #3F51B5
}

skinparam rectangle {
    BackgroundColor #FAFAFA
    BorderColor #455A64
}

skinparam usecase {
    BackgroundColor #EDE7F6
    BorderColor #673AB7
}

left to right direction

actor "Operator Lapangan\\n(Floor Technician)" as Operator
actor "Supervisor Maintenance\\n(Lead Engineer)" as Supervisor
actor "Manajer Pabrik\\n(Plant Manager)" as Manager
actor "Auditor / Asuransi\\n(Third-Party Auditor)" as Auditor

rectangle "Enterprise ERP / SAP" as ERP <<External System>> #E0F2F1
rectangle "Polygon Smart Contract" as Polygon <<External System>> #EDE7F6

rectangle "EchoFactory Platform (System Boundary)" {
    usecase "UC-01: Rekam & Pindai\\nAkustik Mesin" as UC1
    usecase "UC-02: Tanya Jawab Suara\\nHands-Free" as UC2
    usecase "UC-03: Kalibrasi Baseline\\nMesin Baru" as UC3
    usecase "UC-04: Diagnosis Multimodal\\n& SOP RAG" as UC4
    usecase "UC-05: Generate Work Order\\n& Cek Stok Sparepart" as UC5
    usecase "UC-06: Monitoring Dasbor\\nArmada & Estimasi RUL" as UC6
    usecase "UC-07: Verifikasi Paspor\\nKesehatan On-Chain" as UC7
    usecase "UC-08: Eksekusi Klaim\\nGaransi Parametrik" as UC8

    UC1 .> UC4 : <<extend>>\\n[jika anomali]
    UC4 .> UC5 : <<include>>
}

Operator -- UC1
Operator -- UC2
Operator -- UC3

Supervisor -- UC3
Supervisor -- UC4
Supervisor -- UC5
Supervisor -- UC8

Manager -- UC6

Auditor -- UC7
Auditor -- UC8

UC1 -- Polygon
UC3 -- Polygon
UC5 -- ERP
UC7 -- Polygon
UC8 -- Polygon
@enduml
"""

FLOWCHART_PUML = """@startuml
skinparam shadowing false
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam ActivityBackgroundColor #E8EAF6
skinparam ActivityBorderColor #3F51B5
skinparam DiamondBackgroundColor #FFF9C4
skinparam DiamondBorderColor #FBC02D

start
:Mulai Inspeksi Harian Mesin;
:Operator merekam audio mesin 10 detik via HP / Edge Node (16kHz PCM);
:STgram-MFN mengekstrak Mel-Spectrogram & Linear STFT (<50ms);
:KNN-k5 menghitung jarak anomali terhadap baseline normal;

if (Apakah Skor Anomali > Threshold?) then (Ya - Anomali Terdeteksi)
    #FFCDD2:Bunyikan Alarm Merah & Tampilkan Visual Spektrogram;
    :Kirim Dual-Spectrogram & Telemetri ke Gemini Multimodal Core;
    :RAG menelusuri Basis Data Manual Mesin & Standar ISO 10816;
    :Model Degradasi mengestimasi Remaining Useful Life (RUL);
    :Terbitkan Laporan Diagnostik & Rekomendasi Suku Cadang;
    
    if (Supervisor Menyetujui Work Order?) then (Ya)
        #C8E6C9:Panggil API Enterprise ERP / SAP Pabrik;
        :Cek Ketersediaan Stok Sparepart di Gudang Utama;
        :Terbitkan Nomor Tiket WO Resmi (e.g. WO-2026-0814-09);
        :Kirim Notifikasi Tugas Kerja & Part ke Tim Teknisi Shift;
    else (Tidak)
        :Revisi / Jadwalkan Ulang Inspeksi Lapangan;
    endif
else (Tidak - Mesin Sehat)
    #C8E6C9:Tampilkan Kartu Hijau (Pass / Mesin Sehat);
    :Generate Kriptografi SHA-256 Hash Bukti Inspeksi;
    :Simpan Hash Komitmen ke Smart Contract Polygon Amoy;
endif

stop
@enduml
"""

SEQUENCE_PUML = """@startuml
skinparam shadowing false
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam sequenceMessageAlign center

actor "Operator Lapangan" as Op #E3F2FD
actor "Supervisor Maint." as Sup #E8F5E9
boundary "EchoFactory App" as App #EDE7F6
control "STgram-MFN AI" as AI #FFF3E0
control "Gemini & RAG" as GenAI #FCE4EC
entity "SAP / ERP System" as ERP #E0F2F1
database "Polygon Blockchain" as Web3 #EDE7F6

autonumber
Op -> App: Rekam Audio Mesin (10 Detik 16kHz PCM)
App -> AI: Ekstraksi Spektrogram & Hitung Jarak (KNN)
AI --> App: Return Skor Anomali: 0.840 (Alert: Anomali)
App -> Op: Tampilkan Peringatan Visual (Gagal / Rusak)
App -> Sup: Kirim Notifikasi Darurat Anomali Mesin

opt Analisis Akar Masalah & SOP
    Sup -> App: Minta AI Mendiagnosis Akar Masalah
    App -> GenAI: Kirim Gambar Spektrogram + Telemetri Mesin
    GenAI -> GenAI: RAG Standar ISO 10816 & Estimasi RUL Fisik
    GenAI --> App: Hasil: "Bearing Aus. Sisa Umur (RUL): 38 Jam"
    App -> Sup: Tampilkan Laporan Diagnostik AI
end

opt Penerbitan Work Order & Part
    Sup -> App: Approve Work Order (Penggantian Bearing #SKF-6204)
    App -> ERP: Request Pembuatan WO & Cek Stok Sparepart
    ERP -> ERP: Alokasikan Stok Part di Gudang
    ERP --> App: WO Terbit (#WO-2026-0814-09) & Stok Dialokasikan
    App -> Op: Kirim Alert WA (Instruksi Kerja & Batas Waktu)
end

group Integritas Audit On-Chain (Machine Health Passport)
    App -> App: Generate SHA-256 Hash Data Inspeksi Komplit
    App -> Web3: recordInspection(Machine_ID, Score, Status, Hash)
    Web3 --> App: Konfirmasi TxHash Tersimpan di Ledger Polygon
end
@enduml
"""

# ==================== BUILD DOCX ====================

def create_styled_document():
    img_usecase = render_diagram(USECASE_PUML, "plantuml", "usecase_v2.png")
    img_flowchart = render_diagram(FLOWCHART_PUML, "plantuml", "flowchart_v2.png")
    img_sequence = render_diagram(SEQUENCE_PUML, "plantuml", "sequence_v2.png")

    doc = docx.Document()

    # Set Margins
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Styles
    COLOR_PRIMARY = RGBColor(26, 35, 126)   # Deep Navy (#1A237E)
    COLOR_SECONDARY = RGBColor(40, 53, 147) # Slate Blue
    COLOR_TEXT = RGBColor(33, 33, 33)

    def set_font(run, name="Arial", size_pt=11, color=COLOR_TEXT, bold=False, italic=False):
        run.font.name = name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_font(run, size_pt=22, color=COLOR_PRIMARY, bold=True)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        set_font(run, size_pt=12, color=RGBColor(97, 97, 97), italic=True)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_font(run, size_pt=14, color=COLOR_PRIMARY, bold=True)

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_font(run, size_pt=12, color=COLOR_SECONDARY, bold=True)

    def add_p(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            set_font(r_pre, size_pt=10.5, color=COLOR_TEXT, bold=True)
        r_txt = p.add_run(text)
        set_font(r_txt, size_pt=10.5, color=COLOR_TEXT)
        return p

    def add_callout(text, bold_title="Catatan Standar UML: "):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        # Light blue background & left border
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EAF6"/>')
        tcPr.append(shd)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(bold_title)
        set_font(r1, size_pt=10, color=COLOR_PRIMARY, bold=True)
        r2 = p.add_run(text)
        set_font(r2, size_pt=10, color=COLOR_TEXT)

    # --- COVER / HEADER ---
    add_title("EchoFactory System Diagrams & Workflow")
    add_subtitle("Visualisasi UML 2.5: Use Case Diagram, Flowchart Deteksi, dan Sequence Diagram Integritas Data\nCOMPFEST 18 AI Innovation Challenge (Smart Manufacturing)")

    # --- SECTION 1: USE CASE DIAGRAM ---
    add_h1("1. UML Use Case Diagram (Standar Rekayasa Perangkat Lunak UML 2.5)")
    add_p("Diagram Use Case di bawah ini memetakan interaksi fungsional antara seluruh Aktor Pengguna (Manusia di sebelah kiri), Aktor Sistem Eksternal (di sebelah kanan), dan Batas Sistem (System Boundary) dari platform EchoFactory. Seluruh relasi telah diselaraskan dengan kaidah baku UML 2.5:")
    
    # Embed Use Case Image
    p_img1 = doc.add_paragraph()
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img1.paragraph_format.space_before = Pt(6)
    p_img1.paragraph_format.space_after = Pt(8)
    p_img1.add_run().add_picture(img_usecase, width=Inches(6.2))

    add_callout(
        "1. System Boundary Box membatasi fungsionalitas internal EchoFactory dari aktor eksternal.\n"
        "2. Relasi Asosiasi (garis solid) digunakan untuk menghubungkan Aktor primer/sekunder ke Use Case.\n"
        "3. Relasi Dependensi <<extend>> digunakan saat anomali terdeteksi pada UC-01 yang secara kondisional memicu UC-04.\n"
        "4. Relasi Dependensi <<include>> digunakan karena rekomendasi perbaikan dari UC-04 selalu diteruskan ke pembuatan tiket pada UC-05.",
        "Kepatuhan Standar Notasi UML 2.5: "
    )

    # --- TABEL USE CASE MATRIX ---
    add_h2("Tabel Matriks Spesifikasi Aktor & Use Case")
    
    uc_table = doc.add_table(rows=9, cols=4)
    uc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ID", "Nama Use Case", "Aktor Utama / Sekunder", "Deskripsi Fungsional"]
    
    # Header styling
    for col_idx, h_text in enumerate(headers):
        cell = uc_table.cell(0, col_idx)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A237E"/>')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_text)
        set_font(run, size_pt=10, color=RGBColor(255, 255, 255), bold=True)

    rows_data = [
        ("UC-01", "Rekam & Pindai Akustik", "Operator Lapangan, Polygon Smart Contract", "Merekam suara mesin 10 detik dan inferensi ONNX STgram-MFN (<50ms) untuk deteksi Pass/Fail."),
        ("UC-02", "Tanya Jawab Hands-Free", "Operator Lapangan, Gemini AI Core", "Asisten suara interaktif (STT/TTS) berbahasa Indonesia saat tangan operator memegang perkakas."),
        ("UC-03", "Kalibrasi Baseline Mesin", "Supervisor Maint., Operator, Polygon", "Merekam 3 sampel audio normal mesin baru untuk menghitung vektor centroid baseline & minting Genesis Token."),
        ("UC-04", "Diagnosis Multimodal & RAG", "Supervisor Maint., Gemini Flash", "Analisis visual spektrogram + penelusuran manual ISO 10816 + estimasi sisa umur operasional (RUL)."),
        ("UC-05", "Generate Work Order", "Supervisor Maint., Enterprise ERP/SAP", "Menerbitkan tiket perbaikan resmi dan memeriksa ketersediaan stok sparepart (#SKF-6204) di gudang."),
        ("UC-06", "Monitoring Dasbor & RUL", "Manajer Pabrik", "Dasbor armada mesin real-time dengan kode warna risiko kegagalan dan estimasi penghematan downtime."),
        ("UC-07", "Verifikasi Paspor On-Chain", "Auditor / Calon Pembeli, Polygon", "Memindai QR Code mesin untuk memvalidasi keaslian riwayat servis on-chain tanpa perantara."),
        ("UC-08", "Eksekusi Klaim Garansi", "Supervisor, Pihak Asuransi / OEM, Polygon", "Penyelesaian klaim garansi otomatis berbasis konsensus audit kepatuhan inspeksi mesin di smart contract.")
    ]

    for row_idx, data in enumerate(rows_data, start=1):
        for col_idx, val in enumerate(data):
            cell = uc_table.cell(row_idx, col_idx)
            tcPr = cell._tc.get_or_add_tcPr()
            bg_color = "F5F5F5" if row_idx % 2 == 1 else "FFFFFF"
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_font(run, size_pt=9.5, color=COLOR_TEXT)

    # Set column widths
    widths = [Inches(0.8), Inches(1.8), Inches(1.8), Inches(2.1)]
    for row in uc_table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w

    # --- SECTION 2: FLOWCHART ---
    doc.add_page_break()
    add_h1("2. Flowchart Logika Deteksi & Mitigasi Anomali Mesin")
    add_p("Flowchart di bawah ini merinci alur pengambilan keputusan logika sistem sejak Operator melakukan perekaman audio hingga diterbitkannya Work Order di sistem ERP atau commit hash integritas ke blockchain Polygon:")

    p_img2 = doc.add_paragraph()
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.paragraph_format.space_before = Pt(6)
    p_img2.paragraph_format.space_after = Pt(8)
    p_img2.add_run().add_picture(img_flowchart, width=Inches(5.0))

    add_p("1. Jalur Normal (Mesin Sehat):", "• ")
    add_p("Jika skor anomali di bawah ambang batas (Score ≤ Threshold), sistem langsung menampilkan status hijau (Pass), membuat SHA-256 hash, dan menyimpannya secara otomatis ke Smart Contract Polygon Amoy.")
    add_p("2. Jalur Anomali (Kerusakan Terdeteksi):", "• ")
    add_p("Jika skor anomali melebihi ambang batas, alarm berbunyi dan memicu Gemini Flash Multimodal untuk mendiagnosis komponen spesifik yang rusak serta mengestimasi RUL sebelum Work Order diterbitkan.")

    # --- SECTION 3: SEQUENCE DIAGRAM ---
    doc.add_page_break()
    add_h1("3. Sequence Diagram (Workflow Integritas Data & Aksi Enterprise)")
    add_p("Sequence Diagram ini memvisualisasikan urutan waktu (message trace) interaksi antar komponen sistem dalam skenario end-to-end terburuk (Anomali terdeteksi, analisis multimodal, pembuatan tiket perbaikan ERP, hingga pencatatan audit permanen on-chain):")

    p_img3 = doc.add_paragraph()
    p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img3.paragraph_format.space_before = Pt(6)
    p_img3.paragraph_format.space_after = Pt(8)
    p_img3.add_run().add_picture(img_sequence, width=Inches(6.3))

    add_callout(
        "Pencatatan data inspeksi ke blockchain pada Langkah 13-14 menggunakan komitmen kriptografi keccak256(Audio_Hash + Report_Hash + Timestamp) yang ditandatangani secara digital, menjamin data audit kesehatan mesin tidak dapat diubah (tamper-proof) oleh pihak manapun.",
        "Jaminan Integritas Data On-Chain: "
    )

    # Save
    doc.save(TARGET_DOCX)
    print(f"[SUCCESS] Dokumen Word berhasil diperbarui: {TARGET_DOCX}")

if __name__ == "__main__":
    create_styled_document()
