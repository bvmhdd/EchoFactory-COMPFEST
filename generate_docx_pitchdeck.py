import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PITCHDECK WORKSHEET IDEASI GAGASAN\nCOMPFEST 18 AI INNOVATION CHALLENGE (AIC)")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Rangkuman 3 Ide Inovasi Terpilih (JeniusWaste, KoperasiSurya, InnoVault)\nSub-Tema: Smart Manufacturing & AI for Backbone Economy")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph()

    def add_section_header(title, weight=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        if weight:
            r_w = p.add_run(f" ({weight})")
            r_w.font.size = Pt(11)
            r_w.font.bold = True
            r_w.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        return p

    def add_field_content(label, content):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r_l = p.add_run(label + ": ")
        r_l.font.bold = True
        r_l.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        p_ans = doc.add_paragraph()
        p_ans.paragraph_format.left_indent = Inches(0.2)
        p_ans.paragraph_format.space_after = Pt(6)
        r_ans = p_ans.add_run(content)
        r_ans.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        return p_ans

    ideas_data = [
        {
            "num": "1",
            "name": "JeniusWaste — AI Limbah-ke-Energi Terverifikasi & Blockchain Green Certificate",
            "meta": [
                ("Nama Inovasi", "JeniusWaste (AI Waste-to-Energy & Blockchain Certificate)"),
                ("Sub-Tema", "Smart Manufacturing & Sustainability"),
                ("Teknologi Stack", "Fine-tuned XGBoost Regressor + LayoutLMv3 + FastAPI + Next.js + Polygon Amoy Testnet (ERC-1155)")
            ],
            "b1_1": "15.000 industri tahu-tempe dan UMKM olahan pangan menghasilkan jutaan liter limbah cair (COD tinggi) yang 99% dibuang ke sungai. Data LHK mencatat limbah tahu menurunkan kualitas air >60% sungai sentra UMKM. Di sisi lain, Bauran Bioenergi 2025 baru 7.45%, padahal 50L limbah tahu mampu hasilkan 1.8 m3 biogas (hemat Rp 500rb - 2jt/bulan).",
            "b1_2": "Target Pengguna: UMKM olahan pangan & Korporasi pencari ESG credit. Pendekatan Baru: Mengubah polutan menjadi aset digital terverifikasi. AI memvalidasi yield biogas & CO2 offset secara objektif, lalu mendigitalisasikan menjadi Green Certificate (NFT) tamper-proof di Blockchain.",
            "b2_1": "UMKM input volume limbah & suhu -> AI prediksi biogas & CO2 offset -> Mint Green Certificate NFT -> Listed di marketplace ESG.",
            "b2_2": "AI Regressor + Anomaly Detector mutlak diperlukan untuk mencegah greenwashing/klaim palsu limbah berdasarkan hukum termodinamika & data biologis.",
            "b3_1": "Dataset riset biogas BRIN/Kemenperin + data sintetik COD/suhu. Model: XGBoost Regressor + Isolation Forest. Stack: Next.js, FastAPI, Polygon Amoy Testnet.",
            "b4": "FE Single input form + Certificate Viewer. BE API sinkron inference + testnet minting. Dockerized setup.",
            "b6": "Monetisasi: 2.5% fee per transaksi Green Certificate. Etika: Perhitungan CO2 mengacu pada standar IPCC."
        },
        {
            "num": "2",
            "name": "KoperasiSurya — AI-Optimized Solar Cooperative & Transparent Blockchain Profit Sharing",
            "meta": [
                ("Nama Inovasi", "KoperasiSurya (Decentralized Solar Cooperative)"),
                ("Sub-Tema", "Smart Manufacturing & Swasembada Energi"),
                ("Teknologi Stack", "Prophet / LSTM + PuLP Linear Programming + FastAPI + Next.js + Polygon Smart Contract")
            ],
            "b1_1": "Modal PLTS Atap mahal (Rp 15-30jt/kWp) tidak terjangkau UMKM individu. Koperasi komunal gagal karena krisis kepercayaan pengurus. Data: Listrik menyumbang 30-40% biaya operasional UMKM manufaktur. Target Presiden Prabowo: PLTS Desa 13 GW dari 100 GW.",
            "b1_2": "Target: Sentra UMKM & KUB. Pendekatan Baru: Menggabungkan AI prediksi energi & pembagian adil dengan Smart Contract Multi-Sig tanpa pengurus manusia yang bisa korupsi.",
            "b2_1": "UMKM patungan via Smart Contract -> Panel surya dipasang di sentra -> AI optimasi alokasi energi -> Dividen penghematan ditransfer otomatis on-chain.",
            "b2_2": "Cuaca fluktuatif. AI (Prophet/LSTM + PuLP) mutlak untuk memprediksi iradiasi & mengalokasikan daya secara optimal ke setiap mesin UMKM.",
            "b3_1": "Dataset: Open-Meteo API / PVGIS + PLN B1 tariff dataset. Tech Stack: Next.js, FastAPI, Polygon (ERC-20 Solar Token + Treasury Smart Contract).",
            "b4": "FE Dashboard produksi solar & profit sharing. BE API sinkron alokasi token. Dockerized setup.",
            "b6": "Monetisasi: Platform fee 1.5% dari profit sharing bulanan. Transparansi 100% on-chain."
        },
        {
            "num": "3",
            "name": "InnoVault — AI Patent Similarity & Blockchain IP Notarization untuk UMKM",
            "meta": [
                ("Nama Inovasi", "InnoVault (AI Patent Similarity & Prior Art Notarization)"),
                ("Sub-Tema", "Smart Manufacturing & Intellectual Property"),
                ("Teknologi Stack", "Fine-tuned PatentBERT + LLaMA-3 + FAISS + FastAPI + Next.js + Polygon Amoy Testnet")
            ],
            "b1_1": "UMKM manufaktur (furnitur, alat pertanian) buat inovasi tapi tak mampu bayar paten formal (Rp 5-20jt, 2-3 thn). Ide dicuri pabrik besar tanpa bukti awal. Data: <5% dari 64jt UMKM punya KI resmi (DJKI). Kerugian miliaran akibat tiruan barang impor.",
            "b1_2": "Target: UMKM Manufaktur Kreatif & Pengrajin. Pendekatan Baru: Instant Interim IP Protection berbiaya ~0 rupiah. AI hitung Novelty Score vs paten global & Blockchain mengunci Proof of Prior Art berstempel waktu.",
            "b2_1": "UMKM upload deskripsi/foto inovasi -> PatentBERT hitung Novelty Score & klaim teknis -> Dockumen di-hash & committed ke Polygon -> Digital IP Certificate terbit.",
            "b2_2": "Bahasa paten rumit. AI NLP (PatentBERT + FAISS) mutlak untuk vector similarity search pada jutaan dokumen paten global.",
            "b3_1": "Dataset: Google Patents Public Dataset + USPTO Open Data. Model: PatentBERT + LLaMA-3. Stack: Next.js, FastAPI, FAISS, Polygon Amoy Testnet.",
            "b4": "FE Form deskripsi inovasi + Novelty Score Viewer. BE Pipeline sinkron FAISS & smart contract notarization.",
            "b6": "Monetisasi: Freemium model (1 certificate gratis/bulan, $5/cert tambahan). Governance: Mematuhi prinsip etika AI WIPO."
        }
    ]

    for idea in ideas_data:
        add_section_header(f"🚀 IDE #{idea['num']}: {idea['name']}")
        
        meta_table = doc.add_table(rows=len(idea['meta']), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (k, v) in enumerate(idea['meta']):
            cell_k = meta_table.cell(i, 0)
            cell_v = meta_table.cell(i, 1)
            cell_k.width = Inches(2.2)
            cell_v.width = Inches(4.3)
            set_cell_background(cell_k, "F0F4F8")
            set_cell_margins(cell_k, 60, 60, 100, 100)
            set_cell_margins(cell_v, 60, 60, 100, 100)
            
            pk = cell_k.paragraphs[0]
            rk = pk.add_run(k)
            rk.font.bold = True
            rk.font.size = Pt(9.5)
            
            pv = cell_v.paragraphs[0]
            rv = pv.add_run(v)
            rv.font.size = Pt(9.5)
            
        doc.add_paragraph()
        add_field_content("1. Pain Point & Data-Driven Urgency", idea['b1_1'])
        add_field_content("2. Kebaruan (Novelty) & Target Pengguna", idea['b1_2'])
        add_field_content("3. Deskripsi Solusi (Proposed Solution)", idea['b2_1'])
        add_field_content("4. Relevansi Penggunaan AI", idea['b2_2'])
        add_field_content("5. Implementasi AI & Arsitektur", idea['b3_1'])
        add_field_content("6. Kesiapan & Batasan MVP Penyisihan", idea['b4'])
        add_field_content("7. Business Value & Governance", idea['b6'])
        
        doc.add_paragraph()
        
    doc.save("Ideation Pitchdeck COMPFEST AIC.docx")
    print("Successfully generated Ideation Pitchdeck COMPFEST AIC.docx")

if __name__ == "__main__":
    create_document()
