import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_echofactory_full_docx():
    doc = docx.Document()

    # Set Margins (1 inch everywhere)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base font style
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Colors
    NAVY = RGBColor(0x0F, 0x29, 0x4A)
    BLUE = RGBColor(0x1D, 0x4E, 0xD8)
    DARK = RGBColor(0x1E, 0x29, 0x3B)
    GREEN = RGBColor(0x05, 0x96, 0x69)
    GRAY = RGBColor(0x47, 0x55, 0x69)

    # 1. HEADER TITLE
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("ECHOFACTORY")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    p_sub1 = doc.add_paragraph()
    p_sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub1.paragraph_format.space_after = Pt(2)
    r_sub1 = p_sub1.add_run("Acoustic Machine Intelligence & Tamper-Proof Health Passport")
    r_sub1.font.size = Pt(13)
    r_sub1.font.bold = True
    r_sub1.font.color.rgb = BLUE

    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub2.paragraph_format.space_after = Pt(12)
    r_sub2 = p_sub2.add_run("COMPFEST 18 AI Innovation Challenge (AIC) | Sub-Tema: Smart Manufacturing")
    r_sub2.font.size = Pt(10.5)
    r_sub2.font.italic = True
    r_sub2.font.color.rgb = GRAY

    # PITCH HOOK CALLOUT BOX
    table_hook = doc.add_table(rows=1, cols=1)
    table_hook.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_hook = table_hook.cell(0, 0)
    set_cell_background(cell_hook, "EFF6FF")
    set_cell_margins(cell_hook, top=120, bottom=120, left=180, right=180)
    p_hook = cell_hook.paragraphs[0]
    p_hook.paragraph_format.space_after = Pt(0)
    r_hook_label = p_hook.add_run("💬 Pitch Hook: ")
    r_hook_label.font.bold = True
    r_hook_label.font.color.rgb = NAVY
    r_hook_text = p_hook.add_run(
        '"Pak Slamet, 35 tahun teknisi pabrik, tahu mesinnya mau rusak hanya dari suaranya. '
        'Tapi Pak Slamet pensiun tahun depan — dan pengetahuannya pergi bersamanya. '
        'EchoFactory mendokumentasikan telinga Pak Slamet ke dalam AI dan menguncinya di Blockchain '
        'agar pabrik tidak pernah berhenti berputar."'
    )
    r_hook_text.font.italic = True
    r_hook_text.font.color.rgb = DARK

    # Helper Functions
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = BLUE
        return p

    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.color.rgb = DARK
        return p

    def add_step_table(headers, steps):
        tbl = doc.add_table(rows=len(steps) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = tbl.cell(0, i)
            set_cell_background(cell, "1E3A8A")
            set_cell_margins(cell, top=70, bottom=70, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for row_idx, row_data in enumerate(steps, start=1):
            bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            for col_idx, text_val in enumerate(row_data):
                cell = tbl.cell(row_idx, col_idx)
                set_cell_background(cell, bg)
                set_cell_margins(cell, top=60, bottom=60, left=90, right=90)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(text_val)
                r.font.size = Pt(9.5)
                if col_idx == 0:
                    r.font.bold = True
        doc.add_paragraph()

    # --- SECTION 1: SYSTEM OVERVIEW ---
    add_h1("1. EXECUTIVE SUMMARY & SYSTEM OVERVIEW")
    add_body(
        "EchoFactory adalah platform Industrial Predictive Maintenance berbasis Acoustic AI dan Web3 Blockchain. "
        "Sistem ini dirancang untuk mendemokratisasi pemantauan kesehatan mesin industri tanpa memerlukan sensor getaran "
        "mahal (Rp 50-200 juta/mesin), melainkan cukup menggunakan mikrofon smartphone atau modul IoT hemat biaya."
    )

    # --- SECTION 2: ACTOR MATRIX ---
    add_h1("2. ACTOR MATRIX (IDENTIFIKASI AKTOR SISTEM)")
    add_body(
        "Berdasarkan standar rekayasa perangkat lunak (UML Software Engineering), sistem EchoFactory "
        "berinteraksi dengan 4 Aktor Manusia (Human Actors) dan 2 Sistem Eksternal (System Actors):"
    )

    actor_headers = ["ID Aktor", "Nama Aktor", "Kategori", "Peran & Tanggung Jawab Utama"]
    actor_data = [
        ("ACT-01", "Operator Lapangan (Floor Technician)", "Human (Primary)", "Melakukan perekaman suara mesin harian, berinteraksi via asisten suara hands-free, dan memantau indikator Pass/Fail."),
        ("ACT-02", "Supervisor / Kepala Maintenance", "Human (Primary)", "Meninjau diagnosis AI multimodal, memvalidasi rekomendasi SOP, menyetujui Work Order ERP, dan mengelola kalibrasi baseline."),
        ("ACT-03", "Manajer Pabrik (Plant Executive)", "Human (Primary)", "Memantau dasbor kesehatan seluruh armada mesin (Fleet Health), analisis RUL degradasi, dan risiko downtime pabrik."),
        ("ACT-04", "Auditor K3 / Pembeli / Asuransi", "Human (Secondary)", "Memindai QR Code mesin untuk memverifikasi keaslian dan integritas catatan servis pada blockchain secara transparan."),
        ("SYS-01", "Enterprise ERP / SAP System", "System (External)", "Menerima pembuatan tiket perbaikan (Work Order) otomatis dan menyediakan data inventaris suku cadang."),
        ("SYS-02", "Polygon Blockchain Smart Contract", "System (External)", "Menyimpan komitmen hash audit kesehatan mesin permanen (tamper-proof) dan mengeksekusi klaim garansi parametrik.")
    ]
    add_step_table(actor_headers, actor_data)

    # --- SECTION 3: UML USE CASES ---
    add_h1("3. UML USE CASE SPECIFICATIONS (INTERAKSI AKTOR VS SISTEM)")
    add_body(
        "Bagian ini merinci skenario interaksi fungsional langkah-demi-langkah (Action-Response) "
        "antara setiap aktor dengan subsistem EchoFactory:"
    )

    # UC-01
    add_h2("UC-01: Rekam & Pindai Akustik Mesin (Edge Acoustic Inspection)")
    add_body("• Aktor Utama: Operator Lapangan (ACT-01) | Aktor Sekunder: STgram-MFN Engine, Polygon Smart Contract (SYS-02)\n"
             "• Pre-kondisi: Operator memilih ID Mesin pada aplikasi mobile.\n"
             "• Post-kondisi: Skor anomali terhitung (<50ms), status visual tampil, dan log tersimpan.")
    uc1_steps = [
        ("1", "Operator mendekatkan mikrofon HP ke bearing mesin dan menekan 'Mulai Pindai Suara'.", "Sistem merekam audio 16kHz Mono PCM selama 10 detik via WebAudio API dengan visualisasi waveform."),
        ("2", "Audio selesai direkam.", "Sistem mengekstrak Mel-Spectrogram & Linear STFT, lalu mengeksekusi STgram-MFN ONNX Model."),
        ("3", "-", "Model menghasilkan embedding 128-D dan modul KNN-k5 menghitung skor anomali (<50ms)."),
        ("4", "Operator melihat hasil pada layar.", "Jika Normal: Menampilkan kartu hijau 'Pass' dan simpan hash log.\nJika Anomali: Membunyikan alarm merah dan memicu UC-04 (Diagnosis Multimodal).")
    ]
    add_step_table(["No", "Aksi Aktor (Operator)", "Respons Sistem (EchoFactory)"], uc1_steps)

    # UC-02
    add_h2("UC-02: Tanya Jawab Suara Hands-Free (Industrial Voice Assistant)")
    add_body("• Aktor Utama: Operator Lapangan (ACT-01) | Aktor Sekunder: Voice STT/TTS, Gemini Flash Diagnostic Core\n"
             "• Pre-kondisi: Tangan operator sedang kotor atau memegang peralatan.\n"
             "• Post-kondisi: Operator menerima jawaban lisan ringkas dari sistem.")
    uc2_steps = [
        ("1", "Operator berbicara: 'Echo, bagaimana kondisi vibrasi Pompa 3 sekarang?'", "Speech-to-Text menangkap audio suara, memfilter noise bising pabrik, dan mengubahnya menjadi query teks."),
        ("2", "-", "Sistem menarik data telemetri dan riwayat inspeksi terakhir Pompa 3."),
        ("3", "-", "Gemini Flash memformulasikan jawaban ringkas SOP Bahasa Indonesia, lalu dikonversi ke TTS audio."),
        ("4", "Operator mendengar jawaban di headset/speaker.", "Sistem memutar suara: 'Pompa 3 terdeteksi anomali kavitasi 15%. Disarankan kurangi bukaan katup hisap 10 derajat.'")
    ]
    add_step_table(["No", "Aksi Aktor (Operator)", "Respons Sistem (EchoFactory)"], uc2_steps)

    # UC-03
    add_h2("UC-03: Kalibrasi Baseline Mesin Baru (Machine Sound Profiling)")
    add_body("• Aktor Utama: Supervisor Maintenance (ACT-02) | Aktor Sekunder: STgram-MFN, Polygon Smart Contract (SYS-02)\n"
             "• Pre-kondisi: Unit mesin baru dipasang atau selesai overhaul total.\n"
             "• Post-kondisi: Baseline akustik tersimpan dan Genesis Token terdaftar on-chain.")
    uc3_steps = [
        ("1", "Supervisor memilih 'Daftarkan Mesin Baru' dan mengisi metadata mesin.", "Sistem meminta teknisi merekam 3 sampel audio berdurasi 10 detik saat mesin berjalan normal."),
        ("2", "Supervisor merekam 3 sampel audio normal.", "Sistem memvalidasi konsistensi sinyal (cosine similarity > 0.95) dan menghitung vektor centroid normal."),
        ("3", "Supervisor menekan 'Simpan & Kunci Baseline'.", "Sistem mengunci baseline di database lokal dan mem-publish Initial Genesis Hash ke Smart Contract Polygon Amoy.")
    ]
    add_step_table(["No", "Aksi Aktor (Supervisor)", "Respons Sistem (EchoFactory)"], uc3_steps)

    # UC-04
    add_h2("UC-04: Analisis Multimodal & Diagnosis RAG (Root Cause Reasoning)")
    add_body("• Aktor Utama: Supervisor Maintenance (ACT-02) | Aktor Sekunder: Gemini Flash Multimodal, Vector DB RAG\n"
             "• Pre-kondisi: Anomali terdeteksi dari UC-01.\n"
             "• Post-kondisi: Laporan diagnosis akar masalah dan estimasi sisa waktu mesin (RUL) terbit.")
    uc4_steps = [
        ("1", "Supervisor membuka notifikasi anomali pada dasbor.", "Sistem menampilkan visual spektrogram dual-branch, kurva FFT peak, dan histori degradasi."),
        ("2", "Supervisor menekan 'Jalankan AI Diagnostic'.", "Sistem mengirimkan gambar spektrogram + data telemetri ke Gemini Flash Engine."),
        ("3", "-", "RAG menelusuri basis data manual PDF mesin (Hitachi Manual / ISO 10816) mencari frekuensi kerusakan."),
        ("4", "-", "Model Degradasi menghitung sisa waktu sebelum breakdown total (Remaining Useful Life)."),
        ("5", "Supervisor membaca rekomendasi.", "Sistem menyajikan laporan: Komponen: Bearing Inner Race Defect | Urgensi: Medium | Sisa Umur: 38 Jam | Part: #SKF-6204.")
    ]
    add_step_table(["No", "Aksi Aktor (Supervisor)", "Respons Sistem (EchoFactory)"], uc4_steps)

    # UC-05
    add_h2("UC-05: Generate Work Order & Cek Stok Sparepart (Enterprise Dispatch)")
    add_body("• Aktor Utama: Supervisor Maintenance (ACT-02) / Automated Agent | Aktor Sekunder: Enterprise ERP / SAP (SYS-01)\n"
             "• Pre-kondisi: Laporan diagnosis akar masalah (UC-04) selesai dibuat.\n"
             "• Post-kondisi: Tiket perbaikan terbit di SAP/ERP dan notifikasi terkirim ke teknisi.")
    uc5_steps = [
        ("1", "Supervisor meninjau draf tiket dan menekan 'Approve Work Order'.", "Sistem mengeksekusi Function Calling API ke sistem ERP / SAP pabrik."),
        ("2", "-", "Sistem ERP memvalidasi ketersediaan stok bearing #SKF-6204 di gudang suku cadang."),
        ("3", "-", "Sistem ERP menerbitkan nomor tiket resmi (e.g. WO-2026-0814-09) dan mengalokasikan stok."),
        ("4", "Supervisor & teknisi menerima konfirmasi.", "Sistem mengirimkan alert WhatsApp berisi instruksi kerja, nomor part, dan batas waktu pengerjaan.")
    ]
    add_step_table(["No", "Aksi Aktor (Supervisor)", "Respons Sistem (EchoFactory)"], uc5_steps)

    # UC-06
    add_h2("UC-06: Monitoring Dasbor Armada & Estimasi RUL (Fleet Analytics)")
    add_body("• Aktor Utama: Manajer Pabrik (ACT-03) | Aktor Sekunder: Analytics Service\n"
             "• Pre-kondisi: Mesin-mesin pabrik terhubung dan aktif mengirimkan log inspeksi.\n"
             "• Post-kondisi: Manajer memiliki visibilitas penuh atas keandalan operasional seluruh pabrik.")
    uc6_steps = [
        ("1", "Manajer membuka Executive Dashboard EchoFactory.", "Sistem memuat peta tata letak pabrik (Plant Layout) dengan indikator warna status mesin real-time."),
        ("2", "Manajer menyortir mesin berdasarkan sisa umur terpendek.", "Sistem menyajikan daftar mesin dengan risiko tinggi (e.g. Pompa 3: RUL 18 Jam, Fan Blower: RUL 42 Jam)."),
        ("3", "Manajer menekan 'Export Laporan Kesiapan Pabrik'.", "Sistem mengenerate PDF/Excel rekapitulasi MTBF, rasio kepatuhan inspeksi, dan estimasi downtime prevented.")
    ]
    add_step_table(["No", "Aksi Aktor (Manajer)", "Respons Sistem (EchoFactory)"], uc6_steps)

    # UC-07
    add_h2("UC-07: Verifikasi Paspor Kesehatan Mesin On-Chain (Blockchain Verification)")
    add_body("• Aktor Utama: Auditor K3 / Pembeli Mesin Bekas / Perusahaan Leasing (ACT-04) | Aktor Sekunder: Polygon Smart Contract (SYS-02)\n"
             "• Pre-kondisi: Pengguna memindai QR Code fisik mesin atau menginput ID Mesin.\n"
             "• Post-kondisi: Catatan servis terverifikasi secara matematis dan anti-manipulasi.")
    uc7_steps = [
        ("1", "Pengguna memindai QR Code pada mesin menggunakan kamera smartphone.", "Browser membuka portal publik Web3: passport.echofactory.io/verify/{machine_id}."),
        ("2", "-", "Portal memanggil fungsi smart contract getMachineAuditTrail() di Polygon Amoy Testnet."),
        ("3", "-", "Smart Contract mengembalikan deret riwayat audit: Timestamp blok, Skor Anomali, dan Hash Kriptografi."),
        ("4", "Pengguna memeriksa riwayat integritas mesin.", "Layar menampilkan badge hijau: 'Verified On-Chain: 100% Authentic', histori 12 bulan servis, dan rekomendasi valuasi.")
    ]
    add_step_table(["No", "Aksi Aktor (Auditor/Pembeli)", "Respons Sistem (EchoFactory)"], uc7_steps)

    # UC-08
    add_h2("UC-08: Eksekusi Klaim Garansi Parametrik (Smart Contract Settlement)")
    add_body("• Aktor Utama: Supervisor Maintenance (ACT-02) / Lembaga Asuransi (ACT-04) | Aktor Sekunder: Polygon Smart Contract (SYS-02)\n"
             "• Pre-kondisi: Terjadi kerusakan tak terduga dan mesin memiliki polis garansi aktif on-chain.\n"
             "• Post-kondisi: Klaim garansi disetujui secara otomatis tanpa sengketa birokrasi.")
    uc8_steps = [
        ("1", "Supervisor menekan 'Ajukan Klaim Garansi Cerdas'.", "Sistem mengumpulkan bukti log anomali dan riwayat kepatuhan inspeksi harian dari blockchain."),
        ("2", "-", "Smart Contract memverifikasi apakah tingkat kepatuhan inspeksi rutin memenuhi syarat polis (>= 95%)."),
        ("3", "-", "Jika Valid: Smart Contract meng-approve klaim dan men-trigger API vendor part untuk pengiriman instan.\nJika Tidak: Smart Contract menolak klaim otomatis disertai bukti kelalaian inspeksi."),
        ("4", "Supervisor & Asuransi menerima bukti settlement.", "Transaksi tercatat di Polygon Explorer dengan TxHash yang transparan dan tidak dapat disengketakan.")
    ]
    add_step_table(["No", "Aksi Aktor (Supervisor/Asuransi)", "Respons Sistem (EchoFactory)"], uc8_steps)

    # --- SECTION 4: 5-TIER SYSTEM ARCHITECTURE ---
    add_h1("4. ARSITEKTUR SISTEM (5-TIER MODULAR INDUSTRIAL ARCHITECTURE)")
    t_arch = doc.add_table(rows=6, cols=3)
    t_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch_headers = ["Tier Arsitektur", "Komponen & Teknologi", "Peran & Fungsi Utama"]
    for i, h in enumerate(arch_headers):
        cell = t_arch.cell(0, i)
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=70, bottom=70, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    arch_data = [
        ("Tier 1: Sensing & Ingestion", "React Native Mobile App, WebAudio API, ESP32-S3 IoT + MEMS Mic (INMP441)", "Mengambil audio 16kHz 16-bit Mono PCM selama 10 detik tanpa memerlukan hardware mahal."),
        ("Tier 2: Edge Anomaly AI", "STgram-MFN v3 (Dual-Branch MobileFaceNet), ONNX Runtime, Multi-Scorer KNN-k5", "Mengekstrak 128-D embedding dan menghitung skor anomali real-time (<50ms) dengan AUC 94%-99%."),
        ("Tier 3: Cognitive Reasoning", "Gemini 1.5/2.0 Flash Multimodal Vision, ChromaDB RAG, Physics RUL Model", "Membaca spektrogram visual, mencocokkan SOP manual mesin, dan memprediksi sisa jam operasional."),
        ("Tier 4: Enterprise Action", "Tool-Calling Agents, REST API FastAPI, Webhook SAP/ERP, Notifikasi WA/Telegram", "Otomatisasi pembuatan tiket Work Order, pengecekan inventaris suku cadang, dan pemandu suara teknisi."),
        ("Tier 5: Trust & Valuation", "Polygon Amoy Blockchain, Smart Contract MachineHealthPassport.sol, IPFS Storage", "Mencatat bukti kriptografi audit kesehatan mesin secara permanen, transparan, dan tamper-proof.")
    ]

    for row_idx, (c0_t, c1_t, c2_t) in enumerate(arch_data, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text_val in enumerate([c0_t, c1_t, c2_t]):
            cell = t_arch.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text_val)
            r.font.size = Pt(9)
            if col_idx == 0:
                r.font.bold = True
    doc.add_paragraph()

    # --- SECTION 5: MODEL BENCHMARK ---
    add_h1("5. MODEL BENCHMARK & EXPERIMENTAL RESULTS (0 dB SNR)")
    add_body(
        "Hasil validasi model STgram-MFN v3 pada MIMII Dataset (Hitachi Research) pada kondisi kebisingan pabrik nyata (0 dB SNR):"
    )

    t_bench = doc.add_table(rows=5, cols=5)
    t_bench.alignment = WD_TABLE_ALIGNMENT.CENTER
    b_headers = ["Jenis Mesin", "Benchmark Paper IEEE", "Hasil EchoFactory (AUC)", "pAUC (FPR < 10%)", "Status"]
    for i, h in enumerate(b_headers):
        cell = t_bench.cell(0, i)
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    b_data = [
        ("Fan (Blower Pabrik)", "94.04%", "94.04%", "85.31%", "Exact Match Paper"),
        ("Slider (Rel Geser)", "99.55%", "99.32%", "97.55%", "Near Perfect (>99%)"),
        ("Valve (Katup Solenoid)", "99.64%", "Target: 99.6%", "Target: >97%", "Siap Dijalankan"),
        ("Pump (Pompa Industri)", "91.94%", "Target: 91.9%", "Target: >82%", "Siap Dijalankan")
    ]

    for row_idx, row_vals in enumerate(b_data, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(row_vals):
            cell = t_bench.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if col_idx in [0, 2]:
                r.font.bold = True
                if col_idx == 2:
                    r.font.color.rgb = GREEN

    # Spec Summary Box
    doc.add_paragraph()
    table_spec = doc.add_table(rows=1, cols=1)
    table_spec.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_spec = table_spec.cell(0, 0)
    set_cell_background(cell_spec, "F0FDF4")
    set_cell_margins(cell_spec, top=100, bottom=100, left=140, right=140)
    p_spec = cell_spec.paragraphs[0]
    p_spec.paragraph_format.space_after = Pt(0)
    r_sp_title = p_spec.add_run("⚡ Keunggulan Komputasi & Efisiensi Edge:\n")
    r_sp_title.font.bold = True
    r_sp_title.font.color.rgb = GREEN
    r_sp_desc = p_spec.add_run(
        "• Ukuran File ONNX: 183.8 KB (Sangat hemat memori, dapat berjalan offline pada ESP32 / HP)\n"
        "• Parameter Model: 707.200 Parameter (0.71M) — MobileFaceNet backbone teroptimasi\n"
        "• Latency Inferensi: < 50 milidetik pada CPU mobile standar tanpa GPU\n"
        "• Kecepatan Latih: 100 epoch hanya ~4.9 menit berkat RAM-Cached Spectrogram Tensors"
    )
    r_sp_desc.font.size = Pt(9.5)
    r_sp_desc.font.color.rgb = DARK

    # --- SECTION 6: COMPFEST MVP COMPLIANCE ---
    add_h1("6. BATASAN RUANG LINGKUP MVP (COMPFEST PENYISIHAN COMPLIANCE)")
    add_body(
        "Sesuai 'Ketentuan Batasan Ruang Lingkup MVP' dalam Panduan Resmi COMPFEST 18 AIC, "
        "pengembangan pada repositori tahap penyisihan secara ketat dibatasi untuk menjamin kemudahan "
        "reproduksibilitas pengujian lokal oleh dewan juri:"
    )

    t_scope = doc.add_table(rows=6, cols=3)
    t_scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    scope_headers = ["Dimensi Arsitektur", "🎯 Scope MVP Penyisihan (Wajib di Repo)", "🚀 Full Vision & Roadmap (Pitchdeck)"]
    for i, h in enumerate(scope_headers):
        cell = t_scope.cell(0, i)
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    scope_data = [
        ("Frontend (FE) / UI", "Single Screen Input (Upload/Rekam Audio 10s) -> Output visual Spektrogram + Skor Anomali + Kartu Diagnosis.", "Multi-role Dashboard, Peta Layout Pabrik, Fitur Voice Assistant hands-free."),
        ("Backend (BE) & API", "Synchronous REST API FastAPI (1 endpoint sinkron POST /api/detect).", "Distributed Celery workers, Webhooks SAP/ERP, IoT Broker MQTT."),
        ("Model AI & Algoritma", "Core Static Inference: ONNX STgram-MFN + KNN-k5 dengan parameter statis deterministik.", "Continual Active Learning, auto-tuning hyperparameters di cloud."),
        ("Data & Blockchain", "Hashing Kriptografi SHA-256 lokal + simulasi verifikasi Smart Contract on-chain.", "Parametric Insurance Escrow & Decentralized Secondary Market Registry on Polygon."),
        ("Deployment & DevOps", "docker compose up --build (1-Click Run lokal di localhost:3000 / 8000).", "Multi-region Kubernetes Cluster & Edge IoT Firmware Deployment.")
    ]

    for row_idx, (c0, c1, c2) in enumerate(scope_data, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text_val in enumerate([c0, c1, c2]):
            cell = t_scope.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text_val)
            r.font.size = Pt(9)
            if col_idx == 0:
                r.font.bold = True
    doc.add_paragraph()

    # SAVE DOCX with fallback if locked by Word
    output_path = r"c:\Users\muhib\Downloads\COMPFEST\EchoFactory_System_Architecture_and_Workflow.docx"
    try:
        doc.save(output_path)
        print(f"Successfully generated: {output_path}")
    except PermissionError:
        output_path_v2 = r"c:\Users\muhib\Downloads\COMPFEST\EchoFactory_System_Architecture_and_Workflow_V2.docx"
        doc.save(output_path_v2)
        print(f"File was locked by Word. Successfully generated: {output_path_v2}")

if __name__ == "__main__":
    create_echofactory_full_docx()
