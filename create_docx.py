import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("IDEATION PITCHDECK WORKSHEET\nCOMPFEST 18 AI INNOVATION CHALLENGE (AIC)")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Format Pengumpulan Ideation Gagasan Tim - AI for Backbone Economy")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph() # Spacer
    
    # Callout Box
    table_callout = doc.add_table(rows=1, cols=1)
    table_callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table_callout.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p_c = cell.paragraphs[0]
    r_c_bold = p_c.add_run("📌 PANDUAN TIM COMPFEST AIC:\n")
    r_c_bold.font.bold = True
    r_c_bold.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    r_c_text = p_c.add_run(
        "Lembar kerja ini didesain khusus agar setiap anggota tim dapat mengusulkan dan menyusun ideasi gagasan "
        "yang 100% selaras dengan Matriks Penilaian Penyisihan COMPFEST 18 AIC (Bobot: 25% Arsitektur/Teknologi, 20% Orisinalitas/Dampak, "
        "15% Proposal, 15% Kesiapan MVP, 10% Relevansi Tema, 3.5% Business & Governance, 1.5% AIC Talks). "
        "Fokuskan ideasi pada sektor Smart Manufacturing, Smart Logistics, atau Smart Commerce!"
    )
    r_c_text.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Section Header Helper
    def add_section_header(title, weight=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(title)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        if weight:
            r_w = p.add_run(f" ({weight})")
            r_w.font.size = Pt(12)
            r_w.font.bold = True
            r_w.font.color.rgb = RGBColor(0xD9, 0x77, 0x06) # Amber
        return p

    def add_field(label, instruction=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r_l = p.add_run(label + ": ")
        r_l.font.bold = True
        r_l.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        if instruction:
            r_i = p.add_run(instruction)
            r_i.font.italic = True
            r_i.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        p_ans = doc.add_paragraph()
        p_ans.paragraph_format.left_indent = Inches(0.2)
        r_ans = p_ans.add_run("[Isi di sini...]")
        r_ans.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return p_ans

    # METADATA TABLE
    add_section_header("📋 METADATA IDE GAGASAN")
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_headers = [
        ("Nama Tim", "[Isi Nama Tim]"),
        ("Judul / Nama Inovasi", "[Isi Nama produk/Solusi]"),
        ("Sub-Tema Backbone Economy", "[ ] Smart Manufacturing   [ ] Smart Logistics   [ ] Smart Commerce"),
        ("Teknologi Utama / AI Stack", "[Misal: Fine-tuned Llama 3 / YOLOv8 / XGBoost + RAG / FastAPI / React]"),
        ("Kontributor Ide", "[Nama Anggota Tim]")
    ]
    for i, (k, v) in enumerate(meta_headers):
        cell_k = meta_table.cell(i, 0)
        cell_v = meta_table.cell(i, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, "F0F4F8")
        set_cell_margins(cell_k, 80, 80, 100, 100)
        set_cell_margins(cell_v, 80, 80, 100, 100)
        
        pk = cell_k.paragraphs[0]
        rk = pk.add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10)
        
        pv = cell_v.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.size = Pt(10)

    doc.add_paragraph()

    # BAGIAN 1
    add_section_header("📌 BAGIAN 1: ORISINALITAS & DAMPAK SOSIAL", "BOBOT COMPFEST: 20%")
    add_field("1.1 The Pain Point & Data-Driven Urgency", "Jelaskan masalah riil di sektor Backbone Economy yang dipilih beserta minimal 1-2 data statistik kuantitatif riil.")
    add_field("1.2 Kebaruan (Novelty) & Target Pengguna", "Siapa pengguna utama? Mengapa ide ini merupakan pendekatan baru dan apa bedanya dengan solusi yang sudah ada di pasar?")
    add_field("1.3 Skalabilitas & Potensi Global", "Sejauh mana solusi ini dapat diskalakan dan berpotensi memenuhi kebutuhan pasar regional/global?")

    # BAGIAN 2
    add_section_header("🎯 BAGIAN 2: SOLUSI & RELEVANSI TEMA", "BOBOT COMPFEST: 10%")
    add_field("2.1 Deskripsi Solusi (Proposed Solution)", "Uraikan alur kerja produk/layanan yang diusulkan dari sudut pandang pengguna.")
    add_field("2.2 Relevansi Penggunaan AI", "Mengapa AI mutlak diperlukan dan bukan sekadar fitur tempelan? Mengapa solusi ini tidak bisa menggunakan CRUD konvensional?")
    add_field("2.3 Dampak terhadap Backbone Economy", "Bagaimana solusi ini secara khusus meningkatkan efisiensi pada Smart Manufacturing, Logistics, atau Commerce?")

    # BAGIAN 3
    add_section_header("🛠️ BAGIAN 3: IMPLEMENTASI TEKNOLOGI & KEMATANGAN ARSITEKTUR", "BOBOT COMPFEST: 25%")
    add_field("3.1 Alur Dataset (Data Pipeline)", "Sebutkan sumber dataset (publik/sintetik) dan alur preprocessing data sebelum pelatihan model.")
    add_field("3.2 Alur Pengembangan Model AI & Core Inference", "Arsitektur model AI yang digunakan, skema fine-tuning, dan batasan parameter statis saat demo inferensi.")
    add_field("3.3 Arsitektur Sistem Modular (FE - BE - AI)", "Framework Frontend, Backend API, integrasi model server, dan alur komunikasi sinkron.")
    add_field("3.4 Argumentasi Keputusan Teknis (Decision Making)", "Mengapa stack, model, dan arsitektur ini yang dipilih? Berikan argumen teknis berbasis data/efisiensi.")

    # BAGIAN 4
    add_section_header("⚡ BAGIAN 4: KESIAPAN & BATASAN MVP PENYISIHAN", "BOBOT COMPFEST: 15%")
    add_field("4.1 Kedisiplinan Batasan MVP Penyisihan", "Pastikan FE fokus input-output, BE sinkron, AI core inference statis, dan Docker Compose lokal terkonfigurasi.")
    add_field("4.2 Area Pengembangan Babak Final (10-Hour Hackathon Roadmap)", "Komponen atau fitur spesifik apa yang disimpan dan akan dikembangkan secara signifikan pada babak final?")

    # BAGIAN 5
    add_section_header("📄 BAGIAN 5: KUALITAS PROPOSAL & METODOLOGI PENGEMBANGAN", "BOBOT COMPFEST: 15%")
    add_field("5.1 Alur Argumen Proposal 20 Halaman", "Outline singkat bab latar belakang, tujuan, metodologi dataset/model/integrasi, dan kesimpulan.")
    add_field("5.2 Cerita Pengembangan Reflektif & Iteratif", "Tantangan teknis terbesar saat pengembangan dan bagaimana proses iterasi reflektif tim menyelesaikannya.")

    # BAGIAN 6
    add_section_header("⚖️ BAGIAN 6: BUSINESS VALUE & RESPONSIBLE AI", "BONUS COMPFEST: 3.5%")
    add_field("6.1 Model Bisnis & Adopsi Industri", "Skema monetisasi yang realistis dan strategi adopsi oleh pelaku industri di Indonesia.")
    add_field("6.2 Etika, Regulasi & Responsible AI", "Kepatuhan terhadap UU PDP/etika AI, privasi data pengguna, dan keamanan sistem cerdas.")

    doc.add_page_break()

    # SELF EVALUATION MATRIX TABLE
    add_section_header("📊 MATRIKS EVALUASI MANDIRI (SELF-EVALUATION MATRIX) COMPFEST 18 AIC")
    p_mat = doc.add_paragraph("Matriks ini digunakan oleh tim untuk menguji kelayakan ideasi sebelum melangkah ke pembuatan proposal akhir.")
    p_mat.runs[0].font.italic = True
    
    eval_table = doc.add_table(rows=1, cols=5)
    eval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = eval_table.rows[0].cells
    headers = ["Kriteria Penilaian COMPFEST", "Parameter Keberhasilan", "Bobot", "Self-Score (1-5)", "Catatan Perbaikan Tim"]
    widths = [Inches(1.8), Inches(2.2), Inches(0.6), Inches(0.8), Inches(1.1)]
    
    for idx, text in enumerate(headers):
        hdr_cells[idx].width = widths[idx]
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], 100, 100, 80, 80)
        p = hdr_cells[idx].paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ("Implementasi Teknologi & Arsitektur", "Stack proporsional, core inference bersih, arsitektur modular FE-BE-AI terpisah, README & Docker Compose jelas.", "25%", "[ 1-5 ]", ""),
        ("Orisinalitas & Dampak Sosial", "Pendekatan baru & unik, beda dari solusi eksis, relevan dengan masalah individu/bisnis, urgent & ada potensi global.", "20%", "[ 1-5 ]", ""),
        ("Kualitas Proposal & Pengembangan", "Metodologi rinci, alur dataset/model/integrasi logis, decision making berbasis data, cerita iterasi reflektif.", "15%", "[ 1-5 ]", ""),
        ("Kesiapan MVP Penyisihan", "Ruang lingkup MVP tepat (tidak over/underbuilt), fungsionalitas inti teruji, arsitektur fleksibel untuk final.", "15%", "[ 1-5 ]", ""),
        ("Relevansi dengan Tema", "Inovasi sesuai tema AI for Backbone Economy (Smart Manufacturing/Logistics/Commerce), AI relevan.", "10%", "[ 1-5 ]", ""),
        ("[BONUS] Business Value & Governance", "Model bisnis realistis, adopsi industri, pertimbangan regulasi AI, etika & Responsible AI.", "3.5%", "[ 1-5 ]", ""),
        ("[BONUS] AIC Talks", "Anggota tim mengikuti dan mengisi presensi sesi AIC Talks.", "1.5%", "[ 1-5 ]", ""),
        ("TOTAL SKOR EVALUASI", "Target Lolos penyisihan 8 Besar Finalis", "100% (+5%)", "[ Total ]", "Target: Minimal 85+")
    ]

    for r_idx, row in enumerate(rows_data):
        row_cells = eval_table.add_row().cells
        bg_color = "F9FAFB" if r_idx % 2 == 0 else "FFFFFF"
        if r_idx == len(rows_data) - 1:
            bg_color = "FEF3C7" # Yellow total row
        for c_idx, val in enumerate(row):
            row_cells[c_idx].width = widths[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 80, 80, 60, 60)
            p = row_cells[c_idx].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if r_idx == len(rows_data) - 1 or c_idx == 2 or c_idx == 3:
                r.font.bold = True

    doc.add_paragraph()

    # WINNING CHECKLIST
    add_section_header("🏆 CHECKLIST ALASAN KEMENANGAN (WINNING REASON CHECKLIST)")
    checklists = [
        "Data-Driven Problem: Masalah dibuktikan dengan data statistik industri riil (bukan asumsi).",
        "High-Tech & Fine-Tuned Depth: Menggunakan fine-tuning model AI (bukan sekadar wrapper API biasa tanpa preprocessing/RAG).",
        "MVP Scope Discipline: Patuh pada batasan MVP penyisihan COMPFEST (FE Input-Output, BE Synchronous, Dockerized setup).",
        "Clean Modular Architecture: Pembagian peran FE, BE, dan Model Inference Server terpisah rapi.",
        "Backbone Economy Transformation: Memberikan peningkatan efisiensi/produktivitas terukur pada sektor Manufacturing, Logistics, atau Commerce.",
        "Responsible AI & Governance: Menunjukkan kesadaran etika AI, privasi data, dan kepatuhan regulasi di Indonesia."
    ]
    for chk in checklists:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r_box = p.add_run("[  ]  ")
        r_box.font.bold = True
        r_box.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        r_txt = p.add_run(chk)
        r_txt.font.size = Pt(10)

    # Save
    doc.save("Ideation Pitchdeck COMPFEST AIC.docx")
    print("Successfully generated Ideation Pitchdeck COMPFEST AIC.docx")

if __name__ == "__main__":
    create_document()
